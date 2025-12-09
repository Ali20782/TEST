"""
File Processing Module
Handles structured (CSV/XLSX) and unstructured (TXT/DOCX/PDF) files
"""

import pandas as pd
from pathlib import Path
import logging
from typing import Dict, List
import re
import os

logger = logging.getLogger(__name__)

# ============================================================================
# FILE VALIDATION
# ============================================================================

def validate_file_type(filename: str, allowed_extensions: List[str]) -> bool:
    """Validate file extension"""
    extension = Path(filename).suffix.lower().lstrip('.')
    return extension in allowed_extensions

# ============================================================================
# STRUCTURED FILE PROCESSING
# ============================================================================

def process_structured_file(
    file_path: str,
    project_id: int,
    db,
    embedding_service
) -> Dict:
    """
    Process structured event log files (CSV/XLSX)
    Returns: {records_processed: int, embeddings_created: int}
    """
    logger.info(f"Processing structured file: {file_path}")
    
    # Load file
    if file_path.endswith('.csv'):
        df = pd.read_csv(file_path)
    elif file_path.endswith(('.xlsx', '.xls')):
        df = pd.read_excel(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path}")
    
    # Validate required columns
    required_columns = ['case_id', 'activity', 'timestamp']
    missing_columns = [col for col in required_columns if col not in df.columns]
    if missing_columns:
        raise ValueError(f"Missing required columns: {missing_columns}")
    
    # Standardize column names (handle variations)
    column_mapping = {
        'case:concept:name': 'case_id',
        'concept:name': 'activity',
        'time:timestamp': 'timestamp',
        'org:resource': 'resource'
    }
    df = df.rename(columns=column_mapping, errors='ignore')
    
    # Convert timestamp
    df['timestamp'] = pd.to_datetime(df['timestamp'])
    
    # Fill missing optional columns
    if 'resource' not in df.columns:
        df['resource'] = None
    if 'cost' not in df.columns:
        df['cost'] = 0.0
    if 'location' not in df.columns:
        df['location'] = None
    if 'product_type' not in df.columns:
        df['product_type'] = None
    
    logger.info(f"Loaded {len(df)} events from file")
    
    # Insert events into database
    events = df.to_dict('records')
    db.insert_events(project_id, events)
    
    # Get inserted event IDs
    event_ids = db.execute("""
        SELECT id, case_id, activity, timestamp, resource 
        FROM process_mining.event_logs 
        WHERE project_id = %s
        ORDER BY id DESC
        LIMIT %s
    """, (project_id, len(events)), fetch=True)
    
    logger.info(f"Inserted {len(event_ids)} events into database")
    
    # Generate embeddings for events (batch processing for efficiency)
    embeddings_created = 0
    batch_size = 100
    
    for i in range(0, len(event_ids), batch_size):
        batch = event_ids[i:i + batch_size]
        
        # Create text summaries for batch
        event_texts = [create_event_summary(event) for event in batch]
        
        # Generate embeddings in batch
        embeddings = embedding_service.embed_batch(event_texts)
        
        # Store embeddings
        for event_record, embedding in zip(batch, embeddings):
            db.insert_event_embedding(
                project_id=project_id,
                event_id=event_record['id'],
                event_text=create_event_summary(event_record),
                embedding=embedding
            )
            embeddings_created += 1
        
        if (i + batch_size) % 500 == 0:
            logger.info(f"Processed {min(i + batch_size, len(event_ids))}/{len(event_ids)} embeddings")
    
    logger.info(f"Created {embeddings_created} event embeddings")
    
    return {
        'records_processed': len(events),
        'embeddings_created': embeddings_created
    }

def create_event_summary(event: Dict) -> str:
    """Create text summary of event for embedding"""
    parts = [f"Case {event['case_id']}"]
    
    if event.get('activity'):
        parts.append(f"Activity: {event['activity']}")
    
    if event.get('resource'):
        parts.append(f"by {event['resource']}")
    
    if event.get('timestamp'):
        timestamp_str = str(event['timestamp'])
        parts.append(f"at {timestamp_str}")
    
    return " | ".join(parts)

# ============================================================================
# UNSTRUCTURED FILE PROCESSING
# ============================================================================

def process_unstructured_file(
    file_path: str,
    project_id: int,
    db,
    embedding_service
) -> Dict:
    """
    Process unstructured document files (TXT/DOCX/PDF)
    Returns: {chunks_created: int, document_id: int}
    """
    logger.info(f"Processing unstructured file: {file_path}")
    
    # Extract text based on file type
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.txt':
        content_text = extract_text_from_txt(file_path)
    elif file_ext == '.docx':
        content_text = extract_text_from_docx(file_path)
    elif file_ext == '.pdf':
        content_text = extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")
    
    if not content_text or len(content_text.strip()) == 0:
        raise ValueError("No text content extracted from file")
    
    logger.info(f"Extracted {len(content_text)} characters from document")
    
    # Insert document metadata
    filename = Path(file_path).name
    file_size = Path(file_path).stat().st_size
    
    document_id = db.insert_document(
        project_id=project_id,
        filename=filename,
        file_type=file_ext.lstrip('.'),
        file_path=file_path,
        file_size=file_size,
        content_text=content_text
    )
    
    # Chunk the document
    chunks = chunk_text(content_text, chunk_size=500, overlap=100)
    logger.info(f"Created {len(chunks)} chunks")
    
    # Generate embeddings and store (batch processing)
    batch_size = 50
    
    for i in range(0, len(chunks), batch_size):
        batch_chunks = chunks[i:i + batch_size]
        batch_indices = range(i, min(i + batch_size, len(chunks)))
        
        # Generate embeddings in batch
        embeddings = embedding_service.embed_batch(batch_chunks)
        
        # Store chunks with embeddings
        for idx, chunk_text, embedding in zip(batch_indices, batch_chunks, embeddings):
            db.insert_document_chunk(
                document_id=document_id,
                project_id=project_id,
                chunk_index=idx,
                chunk_text=chunk_text,
                embedding=embedding,
                metadata={'chunk_size': len(chunk_text), 'word_count': len(chunk_text.split())}
            )
        
        if (i + batch_size) % 100 == 0:
            logger.info(f"Processed {min(i + batch_size, len(chunks))}/{len(chunks)} chunks")
    
    # Update document chunk count
    db.update_document_chunk_count(document_id, len(chunks))
    
    logger.info(f"Processed document: {len(chunks)} chunks with embeddings")
    
    return {
        'chunks_created': len(chunks),
        'document_id': document_id
    }

# ============================================================================
# TEXT EXTRACTION FUNCTIONS
# ============================================================================

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Failed to extract TXT: {str(e)}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n\n'.join(text)
    except Exception as e:
        logger.error(f"Failed to extract DOCX: {str(e)}")
        raise

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = []
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        return '\n\n'.join(text)
    except Exception as e:
        logger.error(f"Failed to extract PDF: {str(e)}")
        raise

# ============================================================================
# TEXT CHUNKING
# ============================================================================

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
    """
    Split text into overlapping chunks
    
    Args:
        text: Input text
        chunk_size: Target chunk size in words
        overlap: Number of overlapping words between chunks
    
    Returns:
        List of text chunks
    """
    # Clean text - normalize whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    
    if not text:
        return []
    
    # Split into words
    words = text.split()
    
    # If text is shorter than chunk_size, return as single chunk
    if len(words) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(words):
        # Get chunk
        end = start + chunk_size
        chunk_words = words[start:end]
        chunk_text = ' '.join(chunk_words)
        
        # Add chunk
        chunks.append(chunk_text)
        
        # Move start position with overlap
        start = end - overlap
        
        # Prevent infinite loop - if we're near the end, include remaining words
        if start + chunk_size >= len(words):
            # Add final chunk if there are remaining words
            if end < len(words):
                final_chunk = ' '.join(words[end:])
                if len(final_chunk.split()) > overlap // 2:  # Only add if substantial
                    chunks.append(final_chunk)
            break
    
    return chunks

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def get_file_size_mb(file_path: str) -> float:
    """Get file size in megabytes"""
    size_bytes = os.path.getsize(file_path)
    return size_bytes / (1024 * 1024)

def validate_file_size(file_path: str, max_size_mb: int = 100) -> bool:
    """Validate file size is within limits"""
    size_mb = get_file_size_mb(file_path)
    return size_mb <= max_size_mb

def get_row_count(file_path: str) -> int:
    """Get row count for CSV/XLSX files"""
    try:
        if file_path.endswith('.csv'):
            # Fast count for CSV
            with open(file_path, 'r') as f:
                return sum(1 for _ in f) - 1  # Subtract header
        elif file_path.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file_path)
            return len(df)
        else:
            return 0
    except Exception as e:
        logger.error(f"Failed to count rows: {str(e)}")
        return 0

def detect_delimiter(file_path: str, sample_size: int = 1024) -> str:
    """Detect CSV delimiter"""
    import csv
    
    with open(file_path, 'r') as f:
        sample = f.read(sample_size)
        sniffer = csv.Sniffer()
        delimiter = sniffer.sniff(sample).delimiter
        return delimiter

def validate_csv_columns(file_path: str, required_columns: List[str]) -> bool:
    """Validate CSV has required columns"""
    try:
        df = pd.read_csv(file_path, nrows=1)
        return all(col in df.columns for col in required_columns)
    except Exception as e:
        logger.error(f"Failed to validate columns: {str(e)}")
        return False