import pytest
import io
import pandas as pd
from docx import Document

# Note: The 'client' fixture is automatically provided by conftest.py

# --- Health Check Test ---

@pytest.mark.asyncio
async def test_health_check_api(client):
    """Test the /health endpoint to ensure the API is running."""
    response = await client.get("/health")
    # Note: Database status may be "error" in unit tests if env vars aren't set, 
    # but the API status should still be "ok".
    assert response.status_code == 200
    assert response.json()["api_status"] == "ok"

# --- Structured Ingestion Tests ---

@pytest.mark.asyncio
async def test_ingest_structured_csv_success(client):
    """Test successful ingestion of a valid CSV file."""
    # Create a mock CSV file in memory
    csv_content = (
        "case_id,activity,timestamp\n"
        "C001,Start,2025-01-01 10:00:00\n"
        "C001,Activity A,2025-01-01 11:00:00\n"
        "C002,Start,2025-01-02 09:00:00\n"
    )
    
    files = {'file': ('test_log.csv', io.StringIO(csv_content), 'text/csv')}
    response = await client.post("/ingest/structured", files=files)
    
    assert response.status_code == 200
    assert response.json()["metrics"]["total_events"] == 3
    assert response.json()["metrics"]["unique_cases"] == 2
    assert response.json()["status"] == "Structured data accepted for processing."

@pytest.mark.asyncio
async def test_ingest_structured_invalid_file_type(client):
    """Test ingestion with an unsupported file extension."""
    files = {'file': ('document.pdf', io.BytesIO(b'pdf content'), 'application/pdf')}
    response = await client.post("/ingest/structured", files=files)
    
    assert response.status_code == 400
    assert "Invalid file type" in response.json()["detail"]

# --- Unstructured Ingestion Tests ---

@pytest.mark.asyncio
async def test_ingest_unstructured_txt_success(client):
    """Test successful ingestion of a TXT file."""
    txt_content = "This is a process definition document."
    
    files = {'file': ('doc_rules.txt', io.StringIO(txt_content), 'text/plain')}
    response = await client.post("/ingest/unstructured", files=files)
    
    assert response.status_code == 200
    assert response.json()["metrics"]["character_count"] == len(txt_content)
    assert response.json()["status"] == "Unstructured data accepted for processing."

@pytest.mark.asyncio
async def test_ingest_unstructured_docx_success(client):
    """Test successful ingestion of a DOCX file."""
    # Create a mock DOCX file in memory
    document = Document()
    document.add_paragraph('Process Architecture Document.')
    document.add_paragraph('Step 1: Data extraction.')
    
    docx_io = io.BytesIO()
    document.save(docx_io)
    docx_io.seek(0)
    
    files = {'file': ('arch_doc.docx', docx_io, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    response = await client.post("/ingest/unstructured", files=files)

    expected_content = 'Process Architecture Document.\nStep 1: Data extraction.'
    
    assert response.status_code == 200
    assert response.json()["metrics"]["character_count"] > 0
    assert "Process Architecture Document." in response.json()["message"]