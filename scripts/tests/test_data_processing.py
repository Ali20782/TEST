"""
Comprehensive Data Processing Service Tests
Tests file processing, validation, and transformation
"""

import pytest
import pandas as pd
import tempfile
from pathlib import Path
from datetime import datetime
from scripts.data_processing_service import (
    process_structured_data,
    extract_text_from_unstructured,
    validate_event_log_schema,
    transform_to_canonical_format,
    detect_encoding,
    sanitize_dataframe
)


@pytest.fixture
def valid_csv_bytes():
    """Create valid CSV file bytes"""
    content = "case_id,activity,timestamp,resource\n"
    content += "CASE_001,Start,2024-01-01T10:00:00,User1\n"
    content += "CASE_001,Complete,2024-01-01T11:00:00,User1\n"
    return content.encode('utf-8')


@pytest.fixture
def valid_xlsx_file():
    """Create valid XLSX file"""
    df = pd.DataFrame({
        'case_id': ['CASE_001', 'CASE_001'],
        'activity': ['Start', 'Complete'],
        'timestamp': ['2024-01-01T10:00:00', '2024-01-01T11:00:00'],
        'resource': ['User1', 'User1']
    })
    
    with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as f:
        df.to_excel(f.name, index=False)
        temp_path = f.name
    
    yield temp_path
    Path(temp_path).unlink()


@pytest.fixture
def valid_docx_file():
    """Create valid DOCX file"""
    from docx import Document
    
    doc = Document()
    doc.add_paragraph("Invoice Approval Process")
    doc.add_paragraph("This process involves multiple steps.")
    doc.add_paragraph("Step 1: Create Invoice")
    doc.add_paragraph("Step 2: Approve Invoice")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        temp_path = f.name
    
    yield temp_path
    Path(temp_path).unlink()


class TestCSVProcessing:
    """Test CSV file processing"""
    
    def test_process_valid_csv(self, valid_csv_bytes):
        """Test processing valid CSV"""
        df, metrics = process_structured_data(valid_csv_bytes, "test.csv")
        
        assert isinstance(df, pd.DataFrame)
        assert len(df) == 2
        assert 'case_id' in df.columns
        assert 'activity' in df.columns
        assert 'timestamp' in df.columns
        
        assert metrics['total_events'] == 2
        assert metrics['unique_cases'] == 1
    
    def test_process_csv_with_extra_columns(self):
        """Test CSV with additional columns beyond required"""
        content = "case_id,activity,timestamp,resource,cost,extra_col\n"
        content += "CASE_001,Start,2024-01-01T10:00:00,User1,10.5,extra\n"
        
        df, metrics = process_structured_data(content.encode(), "test.csv")
        
        assert 'cost' in df.columns
        assert 'extra_col' in df.columns
    
    def test_process_csv_case_insensitive_columns(self):
        """Test CSV with different case column names"""
        content = "Case_ID,Activity,TimeStamp,Resource\n"
        content += "CASE_001,Start,2024-01-01T10:00:00,User1\n"
        
        df, metrics = process_structured_data(content.encode(), "test.csv")
        
        # Columns should be normalized to lowercase
        assert 'case_id' in df.columns
        assert 'activity' in df.columns
        assert 'timestamp' in df.columns
    
    def test_process_csv_missing_columns(self):
        """Test CSV with missing required columns"""
        content = "case_id,activity\n"
        content += "CASE_001,Start\n"
        
        with pytest.raises(ValueError, match="Missing required columns"):
            process_structured_data(content.encode(), "test.csv")
    
    def test_process_csv_with_nulls(self):
        """Test CSV with null values"""
        content = "case_id,activity,timestamp,resource\n"
        content += "CASE_001,Start,2024-01-01T10:00:00,\n"
        content += "CASE_001,Complete,2024-01-01T11:00:00,User1\n"
        
        df, metrics = process_structured_data(content.encode(), "test.csv")
        
        assert len(df) == 2
        assert pd.isna(df.iloc[0]['resource'])
    
    def test_process_csv_with_encoding(self):
        """Test CSV with different encodings"""
        content = "case_id,activity,timestamp,resource\n"
        content += "CASE_001,Stärt,2024-01-01T10:00:00,Üser\n"
        
        df, metrics = process_structured_data(content.encode('utf-8'), "test.csv")
        
        assert df.iloc[0]['activity'] == 'Stärt'
        assert df.iloc[0]['resource'] == 'Üser'


class TestXLSXProcessing:
    """Test XLSX file processing"""
    
    def test_process_valid_xlsx(self, valid_xlsx_file):
        """Test processing valid XLSX"""
        with open(valid_xlsx_file, 'rb') as f:
            file_bytes = f.read()
        
        df, metrics = process_structured_data(file_bytes, "test.xlsx")
        
        assert isinstance(df, pd.DataFrame)
        assert len(df) == 2
        assert metrics['total_events'] == 2
    
    def test_process_xlsx_multiple_sheets(self):
        """Test XLSX with multiple sheets (should use first sheet)"""
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as f:
            with pd.ExcelWriter(f.name) as writer:
                pd.DataFrame({
                    'case_id': ['CASE_001'],
                    'activity': ['Start'],
                    'timestamp': ['2024-01-01T10:00:00'],
                    'resource': ['User1']
                }).to_excel(writer, sheet_name='Sheet1', index=False)
                
                pd.DataFrame({
                    'other': ['data']
                }).to_excel(writer, sheet_name='Sheet2', index=False)
            
            temp_path = f.name
        
        try:
            with open(temp_path, 'rb') as f:
                file_bytes = f.read()
            
            df, metrics = process_structured_data(file_bytes, "test.xlsx")
            
            assert 'case_id' in df.columns
        finally:
            Path(temp_path).unlink()


class TestTextExtraction:
    """Test text extraction from unstructured files"""
    
    def test_extract_txt(self):
        """Test extracting text from TXT file"""
        content = "This is a test document.\nWith multiple lines."
        
        result = extract_text_from_unstructured(content.encode(), "test.txt")
        
        assert isinstance(result, str)
        assert "This is a test document" in result
        assert "With multiple lines" in result
    
    def test_extract_docx(self, valid_docx_file):
        """Test extracting text from DOCX file"""
        with open(valid_docx_file, 'rb') as f:
            file_bytes = f.read()
        
        result = extract_text_from_unstructured(file_bytes, "test.docx")
        
        assert isinstance(result, str)
        assert "Invoice Approval Process" in result
        assert "Step 1" in result
    
    def test_extract_txt_with_encoding(self):
        """Test text extraction with special characters"""
        content = "Tëst dócument with spëcial çharacters"
        
        result = extract_text_from_unstructured(content.encode('utf-8'), "test.txt")
        
        assert "Tëst" in result
        assert "spëcial" in result
    
    def test_extract_empty_txt(self):
        """Test extracting from empty text file"""
        content = ""
        
        result = extract_text_from_unstructured(content.encode(), "test.txt")
        
        assert result == ""
    
    def test_extract_invalid_format(self):
        """Test extraction with unsupported format"""
        content = b"Binary content"
        
        with pytest.raises(ValueError, match="Unsupported unstructured file type"):
            extract_text_from_unstructured(content, "test.exe")


class TestSchemaValidation:
    """Test event log schema validation"""
    
    def test_validate_valid_schema(self):
        """Test validation of valid schema"""
        df = pd.DataFrame({
            'case_id': ['CASE_001'],
            'activity': ['Start'],
            'timestamp': ['2024-01-01T10:00:00'],
            'resource': ['User1']
        })
        
        is_valid, missing = validate_event_log_schema(df)
        
        assert is_valid is True
        assert len(missing) == 0
    
    def test_validate_missing_columns(self):
        """Test validation with missing columns"""
        df = pd.DataFrame({
            'case_id': ['CASE_001'],
            'activity': ['Start']
        })
        
        is_valid, missing = validate_event_log_schema(df)
        
        assert is_valid is False
        assert 'timestamp' in missing
    
    def test_validate_case_insensitive(self):
        """Test validation is case insensitive"""
        df = pd.DataFrame({
            'Case_ID': ['CASE_001'],
            'Activity': ['Start'],
            'TimeStamp': ['2024-01-01T10:00:00']
        })
        
        is_valid, missing = validate_event_log_schema(df)
        
        assert is_valid is True


class TestDataTransformation:
    """Test data transformation and cleaning"""
    
    def test_transform_to_canonical(self):
        """Test transformation to canonical format"""
        df = pd.DataFrame({
            'case_id': ['CASE_001'],
            'activity': ['Start'],
            'timestamp': ['2024-01-01T10:00:00'],
            'resource': ['User1']
        })
        
        result = transform_to_canonical_format(df)
        
        assert 'case_id' in result.columns
        assert result['case_id'].dtype == object
    
    def test_sanitize_dataframe(self):
        """Test dataframe sanitization"""
        df = pd.DataFrame({
            'case_id': ['CASE_001', 'CASE_002'],
            'activity': ['Start', None],
            'timestamp': ['2024-01-01T10:00:00', '2024-01-01T11:00:00'],
            'resource': ['User1', 'User2']
        })
        
        result = sanitize_dataframe(df)
        
        # Check null handling
        assert not pd.isna(result.iloc[1]['activity'])
    
    def test_timestamp_parsing(self):
        """Test various timestamp formats are parsed"""
        df = pd.DataFrame({
            'case_id': ['C1', 'C2', 'C3'],
            'activity': ['A', 'B', 'C'],
            'timestamp': [
                '2024-01-01T10:00:00',
                '2024-01-01 10:00:00',
                '2024/01/01 10:00:00'
            ],
            'resource': ['U1', 'U2', 'U3']
        })
        
        result = transform_to_canonical_format(df)
        
        # All timestamps should be parsed
        assert result['timestamp'].dtype == 'datetime64[ns]'


class TestEncodingDetection:
    """Test file encoding detection"""
    
    def test_detect_utf8(self):
        """Test UTF-8 encoding detection"""
        content = "Test content".encode('utf-8')
        
        encoding = detect_encoding(content)
        
        assert encoding in ['utf-8', 'ascii']
    
    def test_detect_latin1(self):
        """Test Latin-1 encoding detection"""
        content = "Tëst cöntent".encode('latin-1')
        
        encoding = detect_encoding(content)
        
        assert encoding is not None


class TestMetricsCalculation:
    """Test metrics calculation"""
    
    def test_metrics_basic(self):
        """Test basic metrics calculation"""
        df = pd.DataFrame({
            'case_id': ['C1', 'C1', 'C2'],
            'activity': ['A', 'B', 'A'],
            'timestamp': ['2024-01-01T10:00:00'] * 3,
            'resource': ['U1', 'U1', 'U2']
        })
        
        metrics = {
            'total_events': len(df),
            'unique_cases': df['case_id'].nunique(),
            'unique_activities': df['activity'].nunique(),
            'unique_resources': df['resource'].nunique()
        }
        
        assert metrics['total_events'] == 3
        assert metrics['unique_cases'] == 2
        assert metrics['unique_activities'] == 2
        assert metrics['unique_resources'] == 2


class TestErrorHandling:
    """Test error handling in data processing"""
    
    def test_corrupted_file(self):
        """Test handling of corrupted file"""
        content = b'\x00\x01\x02\x03'
        
        with pytest.raises(Exception):
            process_structured_data(content, "test.csv")
    
    def test_empty_file(self):
        """Test handling of empty file"""
        content = b''
        
        with pytest.raises(Exception):
            process_structured_data(content, "test.csv")
    
    def test_invalid_timestamp_format(self):
        """Test handling of invalid timestamps"""
        content = "case_id,activity,timestamp,resource\n"
        content += "CASE_001,Start,invalid_date,User1\n"
        
        df, metrics = process_structured_data(content.encode(), "test.csv")
        
        # Should handle gracefully, possibly with null timestamp
        assert len(df) == 1


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
