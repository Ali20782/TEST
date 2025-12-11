import pytest
from starlette.testclient import TestClient
from fastapi import FastAPI
from src.api.ingestion_router import router # Import the router we just created
from io import BytesIO
from docx import Document

# --- FastAPI App Fixture ---
# Assuming your main app file is accessible
@pytest.fixture(scope="module")
def app():
    """Defines the FastAPI application for testing."""
    app = FastAPI()
    app.include_router(router)
    return app

@pytest.fixture(scope="module")
def client(app):
    """Provides a TestClient for making requests to the FastAPI app."""
    with TestClient(app) as c:
        yield c

@pytest.fixture(scope="session")
def project_id():
    """Standard project ID fixture."""
    return "TEST_PROJECT_ID_42"

# --- Mock Data Fixtures for File Uploads ---

@pytest.fixture(scope="session")
def mock_csv_file():
    """Creates a mock in-memory CSV file (Structured Data)."""
    # Required columns: case_id, activity, timestamp
    content = (
        "case_id,activity,timestamp,resource\n"
        "C1,Start,2025-01-01 10:00:00,UserA\n"
        "C1,TaskA,2025-01-01 10:30:00,UserA\n"
        "C1,End,2025-01-01 11:00:00,UserB\n"
        "C2,Start,2025-01-02 12:00:00,UserC\n"
    ).encode('utf-8')
    return ("valid_log.csv", BytesIO(content), "text/csv")

@pytest.fixture(scope="session")
def mock_invalid_csv_file():
    """Creates a mock CSV file missing required columns."""
    content = (
        "id,event_name,time\n" # Missing case_id, activity, timestamp
        "1,Start,2025-01-01 10:00:00\n"
    ).encode('utf-8')
    return ("invalid_log.csv", BytesIO(content), "text/csv")


@pytest.fixture(scope="session")
def mock_txt_file():
    """Creates a mock in-memory TXT file (Unstructured Data)."""
    content = "This is a document about process variant analysis. It is unstructured data for RAG."
    return ("valid_doc.txt", BytesIO(content.encode('utf-8')), "text/plain")

@pytest.fixture(scope="session")
def mock_docx_file():
    """Creates a mock in-memory DOCX file (Unstructured Data)."""
    # Use the docx library to create a minimal valid DOCX file
    document = Document()
    document.add_paragraph('This is the first paragraph of the DOCX document.')
    document.add_paragraph('The second paragraph contains keywords like pgvector and fastapi.')
    
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return ("valid_doc.docx", file_stream, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")